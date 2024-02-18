import json
from docx import Document
from docx.shared import RGBColor


def process_color(docx: str):
    """
    处理docx文档的颜色
    :param docx: docx文档的路径
    :return:
    """
    document = Document(docx)
    # 读入颜色
    with open('colors.json') as f:
        colors = json.load(f)
    color_idx = 0

    # 遍历docx的每一行，修改字体颜色
    for pid, paragraph in enumerate(document.paragraphs):
        # 初始样式
        original_run = paragraph.runs[0]
        # 遍历每个字
        for c in paragraph.text:
            run = paragraph.add_run(c)
            # 取出每个字对应的颜色
            color = colors[color_idx]
            color_idx += 1
            # 设置颜色
            run.font.color.rgb = RGBColor(color[0], color[1], color[2])
            # 字体和大小不变
            run.font.name = original_run.font.name
            run.font.size = original_run.font.size
        # 清楚样式
        paragraph.runs[0].clear()
    # 保存
    document.save(f'output_{docx}')
